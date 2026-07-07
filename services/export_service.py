"""Export service: DOCX, PDF, TXT generation."""
import os
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


def _get_all_content(db, project_id, options):
    """Fetch all book content from the database."""
    project = db.execute('SELECT * FROM projects WHERE id = ?', (project_id,)).fetchone()
    chapters = db.execute(
        'SELECT * FROM chapters WHERE project_id = ? AND status = "generated" ORDER BY chapter_number',
        (project_id,)
    ).fetchall()
    front = {}
    if options.get('include_front_matter'):
        rows = db.execute('SELECT type, content FROM front_matter WHERE project_id = ?', (project_id,)).fetchall()
        front = {r['type']: r['content'] for r in rows if r['content']}
    back = {}
    if options.get('include_back_matter'):
        rows = db.execute('SELECT type, content FROM back_matter WHERE project_id = ?', (project_id,)).fetchall()
        back = {r['type']: r['content'] for r in rows if r['content']}
    return project, chapters, front, back


def export_txt(db, project_id, options=None):
    """Export the book as plain text. Returns bytes."""
    if options is None:
        options = {}
    project, chapters, front, back = _get_all_content(db, project_id, options)
    lines = []

    # Title page
    lines.append(project['title'].upper())
    if project['subtitle']:
        lines.append(project['subtitle'])
    lines.append('')
    lines.append(f"Genre: {project['genre']} | Language: {project['language']}")
    lines.append('=' * 60)
    lines.append('')

    # Front matter
    fm_order = ['dedication', 'foreword', 'preface', 'acknowledgements', 'introduction']
    for fm_type in fm_order:
        if fm_type in front:
            lines.append(fm_type.upper().replace('_', ' '))
            lines.append('-' * 40)
            lines.append(front[fm_type])
            lines.append('')

    # Table of contents
    if options.get('include_toc') and chapters:
        lines.append('TABLE OF CONTENTS')
        lines.append('-' * 40)
        for ch in chapters:
            lines.append(f"Chapter {ch['chapter_number']}: {ch['chapter_title']}")
        lines.append('')

    # Chapters
    for ch in chapters:
        lines.append('')
        lines.append(f"CHAPTER {ch['chapter_number']}: {ch['chapter_title'].upper()}")
        lines.append('-' * 60)
        lines.append(ch['content'] or '')
        lines.append('')

    # Back matter
    bm_order = ['conclusion', 'epilogue', 'afterword', 'about_author', 'glossary', 'appendix', 'references']
    for bm_type in bm_order:
        if bm_type in back:
            lines.append(bm_type.upper().replace('_', ' '))
            lines.append('-' * 40)
            lines.append(back[bm_type])
            lines.append('')

    return '\n'.join(lines).encode('utf-8')


def export_docx(db, project_id, options=None):
    """Export the book as a DOCX file. Returns bytes."""
    if options is None:
        options = {}
    project, chapters, front, back = _get_all_content(db, project_id, options)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project['title'])
    run.bold = True
    run.font.size = Pt(28)

    if project['subtitle']:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(project['subtitle'])
        run.font.size = Pt(16)
        run.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Genre: {project['genre']} | Language: {project['language']}")

    if options.get('include_front_matter') and any(front.values()):
        doc.add_page_break()
        fm_order = ['dedication', 'foreword', 'preface', 'acknowledgements', 'introduction']
        for fm_type in fm_order:
            if fm_type in front:
                heading = doc.add_heading(fm_type.replace('_', ' ').title(), level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for para_text in front[fm_type].split('\n'):
                    p = doc.add_paragraph(para_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if options.get('include_toc') and chapters:
        doc.add_page_break()
        toc_heading = doc.add_heading('Table of Contents', level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ch in chapters:
            p = doc.add_paragraph(f"Chapter {ch['chapter_number']}: {ch['chapter_title']}")
            p.paragraph_format.left_indent = Inches(0.25)

    # Chapters
    for ch in chapters:
        doc.add_page_break()
        heading = doc.add_heading(f"Chapter {ch['chapter_number']}: {ch['chapter_title']}", level=1)
        if options.get('include_headers'):
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        content = ch['content'] or ''
        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(6)

    # Back matter
    if options.get('include_back_matter') and any(back.values()):
        bm_order = ['conclusion', 'epilogue', 'afterword', 'about_author', 'glossary', 'appendix', 'references']
        for bm_type in bm_order:
            if bm_type in back:
                doc.add_page_break()
                heading = doc.add_heading(bm_type.replace('_', ' ').title(), level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for para_text in back[bm_type].split('\n'):
                    p = doc.add_paragraph(para_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_pdf(db, project_id, options=None):
    """Export the book as a PDF. Returns bytes."""
    if options is None:
        options = {}
    project, chapters, front, back = _get_all_content(db, project_id, options)

    buf = io.BytesIO()
    title = project['title']
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=1.25 * inch,
        leftMargin=1.25 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        title=title,
        author='KDP Novel & Storybook Creator',
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('BookTitle', parent=styles['Title'],
                                  fontSize=28, spaceAfter=12, alignment=TA_CENTER)
    subtitle_style = ParagraphStyle('BookSubtitle', parent=styles['Normal'],
                                     fontSize=16, italic=True, spaceAfter=8, alignment=TA_CENTER)
    chapter_title_style = ParagraphStyle('ChapterTitle', parent=styles['Heading1'],
                                          fontSize=18, spaceAfter=18, spaceBefore=12, alignment=TA_CENTER)
    section_title_style = ParagraphStyle('SectionTitle', parent=styles['Heading2'],
                                          fontSize=14, spaceAfter=10, spaceBefore=8, alignment=TA_CENTER)
    body_style = ParagraphStyle('BookBody', parent=styles['Normal'],
                                 fontSize=11, leading=18, spaceAfter=8, alignment=TA_JUSTIFY,
                                 firstLineIndent=0.25 * inch)
    toc_style = ParagraphStyle('TOC', parent=styles['Normal'],
                                fontSize=11, leading=16, spaceAfter=4, leftIndent=0.25 * inch)

    story = []

    # Title page
    story.append(Spacer(1, 1.5 * inch))
    story.append(Paragraph(title, title_style))
    if project['subtitle']:
        story.append(Paragraph(project['subtitle'], subtitle_style))
    story.append(Spacer(1, 0.3 * inch))
    story.append(HRFlowable(width='100%', thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.2 * inch))
    meta_style = ParagraphStyle('Meta', parent=styles['Normal'], fontSize=10,
                                 alignment=TA_CENTER, textColor=colors.grey)
    story.append(Paragraph(f"Genre: {project['genre']} | Language: {project['language']}", meta_style))

    # Front matter
    if options.get('include_front_matter'):
        fm_order = ['dedication', 'foreword', 'preface', 'acknowledgements', 'introduction']
        for fm_type in fm_order:
            if fm_type in front:
                story.append(PageBreak())
                story.append(Paragraph(fm_type.replace('_', ' ').title(), section_title_style))
                story.append(HRFlowable(width='60%', thickness=1, color=colors.grey))
                story.append(Spacer(1, 0.2 * inch))
                for line in front[fm_type].split('\n'):
                    if line.strip():
                        story.append(Paragraph(line.strip(), body_style))

    # TOC
    if options.get('include_toc') and chapters:
        story.append(PageBreak())
        story.append(Paragraph('Table of Contents', section_title_style))
        story.append(HRFlowable(width='60%', thickness=1, color=colors.grey))
        story.append(Spacer(1, 0.2 * inch))
        for ch in chapters:
            story.append(Paragraph(f"Chapter {ch['chapter_number']}: {ch['chapter_title']}", toc_style))

    # Chapters
    for ch in chapters:
        story.append(PageBreak())
        story.append(Paragraph(f"Chapter {ch['chapter_number']}", section_title_style))
        story.append(Paragraph(ch['chapter_title'], chapter_title_style))
        story.append(HRFlowable(width='40%', thickness=0.5, color=colors.grey))
        story.append(Spacer(1, 0.3 * inch))
        content = ch['content'] or ''
        for line in content.split('\n'):
            line = line.strip()
            if line:
                story.append(Paragraph(line, body_style))

    # Back matter
    if options.get('include_back_matter'):
        bm_order = ['conclusion', 'epilogue', 'afterword', 'about_author', 'glossary', 'appendix', 'references']
        for bm_type in bm_order:
            if bm_type in back:
                story.append(PageBreak())
                story.append(Paragraph(bm_type.replace('_', ' ').title(), section_title_style))
                story.append(HRFlowable(width='60%', thickness=1, color=colors.grey))
                story.append(Spacer(1, 0.2 * inch))
                for line in back[bm_type].split('\n'):
                    if line.strip():
                        story.append(Paragraph(line.strip(), body_style))

    def add_page_number(canvas, doc):
        if options.get('include_page_numbers'):
            canvas.saveState()
            canvas.setFont('Helvetica', 9)
            canvas.setFillColor(colors.grey)
            page_num = canvas.getPageNumber()
            canvas.drawCentredString(letter[0] / 2, 0.5 * inch, str(page_num))
            canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    buf.seek(0)
    return buf.read()
